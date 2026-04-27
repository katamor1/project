#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
IBM-Bob JSON -> DOCX generator.

Purpose:
  Generate Word design documents from structured JSON and an optional DOCX template.
  This script intentionally avoids direct hand-written OOXML. It uses python-docx
  and validates unresolved placeholders before returning success.

Supported placeholders:
  {{document.document_name}}
  {{$.document.document_name}}
  {{json.document.document_name}}
  {{item.id}} inside repeated table rows

Supported table repeat marker:
  [[TABLE:traceability.requirements]]
  The row after the marker row is treated as the template row and repeated for each item.

Fallback behavior:
  If no template is supplied, or --fallback is selected, a structured DOCX is generated
  from the JSON keys with IBM-Bob friendly headings and tables.
"""
from __future__ import annotations

import argparse
import copy
import json
import os
import re
import sys
import zipfile
from datetime import date
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

PLACEHOLDER_RE = re.compile(r"\{\{\s*([^{}]+?)\s*\}\}")
TABLE_MARKER_RE = re.compile(r"\[\[TABLE\s*:\s*([^\]]+?)\s*\]\]")
UNRESOLVED_RE = re.compile(r"(\{\{[^{}]+\}\}|\[\[TABLE\s*:[^\]]+\]\])")


def load_json(path: str | Path) -> dict[str, Any]:
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if not isinstance(data, dict):
        raise ValueError("Top-level JSON must be an object/dict.")
    return data


def normalize_path(expr: str) -> str:
    expr = expr.strip()
    if expr.startswith("json."):
        expr = expr[5:]
    if expr.startswith("$."):
        expr = expr[2:]
    if expr == "$":
        return ""
    return expr.strip()


def get_path(data: Any, expr: str, default: Any = "") -> Any:
    expr = normalize_path(expr)
    if expr == "":
        return data
    cur = data
    for raw_part in expr.split("."):
        part = raw_part.strip()
        if part == "":
            continue
        # list index support: sections[0] or 0
        m = re.fullmatch(r"(.+?)\[(\d+)\]", part)
        if m:
            name, idx_s = m.groups()
            cur = get_path(cur, name, default)
            idx = int(idx_s)
            if isinstance(cur, list) and 0 <= idx < len(cur):
                cur = cur[idx]
            else:
                return default
            continue
        if isinstance(cur, dict):
            if part in cur:
                cur = cur[part]
            else:
                return default
        elif isinstance(cur, list):
            if part.isdigit() and int(part) < len(cur):
                cur = cur[int(part)]
            else:
                return default
        else:
            return default
    return cur


def stringify(value: Any, *, max_depth: int = 3, indent: int = 0) -> str:
    """Convert a JSON value to readable text for placeholder replacement."""
    prefix = "  " * indent
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, str):
        return value
    if isinstance(value, list):
        if not value:
            return "（該当なし）"
        lines: list[str] = []
        for item in value:
            if isinstance(item, (str, int, float, bool)) or item is None:
                lines.append(f"{prefix}・{stringify(item, max_depth=max_depth, indent=indent)}")
            elif isinstance(item, dict):
                if max_depth <= 0:
                    lines.append(prefix + json.dumps(item, ensure_ascii=False))
                else:
                    flat = "; ".join(f"{k}: {stringify(v, max_depth=max_depth-1, indent=indent+1)}" for k, v in item.items())
                    lines.append(f"{prefix}・{flat}")
            else:
                lines.append(prefix + json.dumps(item, ensure_ascii=False))
        return "\n".join(lines)
    if isinstance(value, dict):
        if not value:
            return "（該当なし）"
        lines = []
        for k, v in value.items():
            if isinstance(v, (dict, list)) and max_depth > 0:
                sub = stringify(v, max_depth=max_depth-1, indent=indent+1)
                lines.append(f"{prefix}{k}:\n{sub}")
            else:
                lines.append(f"{prefix}{k}: {stringify(v, max_depth=max_depth-1, indent=indent+1)}")
        return "\n".join(lines)
    return str(value)


def set_cell_text(cell, text: str) -> None:
    cell.text = ""
    lines = str(text).splitlines() or [""]
    p = cell.paragraphs[0]
    p.text = lines[0]
    for line in lines[1:]:
        p = cell.add_paragraph(line)


def paragraph_full_text(paragraph) -> str:
    return "".join(run.text for run in paragraph.runs)


def replace_placeholders_in_text(text: str, data: dict[str, Any], item: Any = None) -> str:
    def repl(match: re.Match[str]) -> str:
        expr = match.group(1).strip()
        if expr.startswith("item."):
            if item is None:
                return ""
            return stringify(get_path(item, expr[5:], default=""))
        return stringify(get_path(data, expr, default=""))
    return PLACEHOLDER_RE.sub(repl, text)


def replace_in_paragraph(paragraph, data: dict[str, Any], item: Any = None) -> None:
    text = paragraph_full_text(paragraph)
    if "{{" not in text:
        return
    replaced = replace_placeholders_in_text(text, data, item=item)
    # Preserve paragraph style, but simplify runs to avoid split-placeholder bugs.
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = replaced
    else:
        paragraph.add_run(replaced)


def replace_in_table(table, data: dict[str, Any]) -> None:
    expand_repeat_tables(table, data)
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph(p, data)
            for nested in cell.tables:
                replace_in_table(nested, data)


def row_text(row) -> str:
    return "\n".join(cell.text for cell in row.cells)


def clone_row_after(table, template_row, after_index: int):
    tr = copy.deepcopy(template_row._tr)
    table._tbl.insert(after_index + 1, tr)
    return table.rows[after_index + 1]


def delete_row(table, row_index: int) -> None:
    tr = table.rows[row_index]._tr
    tr.getparent().remove(tr)


def replace_placeholders_in_xml_element(element, data: dict[str, Any], item: Any = None) -> None:
    """Replace placeholders in all w:t nodes under an OOXML element.

    This is used for cloned table rows because python-docx row collections may not
    refresh predictably after low-level row insertion. It is safer to patch the
    cloned XML before inserting it.
    """
    for node in element.xpath(".//w:t"):
        if node.text and "{{" in node.text:
            node.text = replace_placeholders_in_text(node.text, data, item=item)


def expand_repeat_tables(table, data: dict[str, Any]) -> None:
    marker_index = None
    marker_path = None
    rows_snapshot = list(table.rows)
    for i, row in enumerate(rows_snapshot):
        txt = row_text(row)
        m = TABLE_MARKER_RE.search(txt)
        if m:
            marker_index = i
            marker_path = m.group(1).strip()
            break
    if marker_index is None or marker_path is None:
        return

    items = get_path(data, marker_path, default=[])
    if not isinstance(items, list):
        items = []
    template_index = marker_index + 1
    if template_index >= len(rows_snapshot):
        # No template row. Leave marker visible; validator will fail in strict mode.
        return

    marker_tr = rows_snapshot[marker_index]._tr
    template_tr = rows_snapshot[template_index]._tr
    insert_after = template_tr

    for item in items:
        new_tr = copy.deepcopy(template_tr)
        replace_placeholders_in_xml_element(new_tr, data, item=item)
        insert_after.addnext(new_tr)
        insert_after = new_tr

    # Remove original template and marker rows.
    template_tr.getparent().remove(template_tr)
    marker_tr.getparent().remove(marker_tr)


def apply_template(data: dict[str, Any], template_path: str | Path, out_path: str | Path) -> None:
    doc = Document(str(template_path))
    for p in doc.paragraphs:
        replace_in_paragraph(p, data)
    for table in doc.tables:
        replace_in_table(table, data)
    for section in doc.sections:
        for p in section.header.paragraphs:
            replace_in_paragraph(p, data)
        for p in section.footer.paragraphs:
            replace_in_paragraph(p, data)
        for table in section.header.tables:
            replace_in_table(table, data)
        for table in section.footer.tables:
            replace_in_table(table, data)
    doc.save(str(out_path))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_styles(doc: DocxDocument) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Yu Gothic"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    styles["Normal"].font.size = Pt(10.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            styles[style_name].font.name = "Yu Gothic"
            styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    if "Heading 1" in styles:
        styles["Heading 1"].font.size = Pt(16)
    if "Heading 2" in styles:
        styles["Heading 2"].font.size = Pt(13)


def add_heading(doc: DocxDocument, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_kv_table(doc: DocxDocument, rows: list[tuple[str, Any]], title: str | None = None) -> None:
    if title:
        doc.add_paragraph(title, style=None).runs[0].bold = True
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "項目"
    hdr[1].text = "内容"
    set_repeat_table_header(table.rows[0])
    for cell in hdr:
        set_cell_shading(cell, "D9EAF7")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for k, v in rows:
        row = table.add_row().cells
        set_cell_text(row[0], str(k))
        set_cell_text(row[1], stringify(v))
        row[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        row[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph("")


def add_list(doc: DocxDocument, items: Iterable[Any], bullet: bool = True) -> None:
    for item in items:
        if isinstance(item, dict):
            text = "; ".join(f"{k}: {stringify(v)}" for k, v in item.items())
        else:
            text = stringify(item)
        doc.add_paragraph(text, style="List Bullet" if bullet else None)


def add_object_section(doc: DocxDocument, key: str, value: Any, level: int = 1) -> None:
    title = humanize_key(key)
    add_heading(doc, title, min(level, 3))
    if isinstance(value, dict):
        scalar_rows: list[tuple[str, Any]] = []
        complex_items: list[tuple[str, Any]] = []
        for k, v in value.items():
            if isinstance(v, (dict, list)):
                complex_items.append((k, v))
            else:
                scalar_rows.append((humanize_key(k), v))
        if scalar_rows:
            add_kv_table(doc, scalar_rows)
        for k, v in complex_items:
            if isinstance(v, list):
                add_heading(doc, humanize_key(k), min(level + 1, 3))
                add_list_or_table(doc, v)
            else:
                add_object_section(doc, k, v, level + 1)
    elif isinstance(value, list):
        add_list_or_table(doc, value)
    else:
        doc.add_paragraph(stringify(value))


def add_list_or_table(doc: DocxDocument, value: list[Any]) -> None:
    if not value:
        doc.add_paragraph("（該当なし）")
        return
    if all(isinstance(x, dict) for x in value):
        keys: list[str] = []
        for item in value:
            for k in item.keys():
                if k not in keys:
                    keys.append(k)
        # Avoid extremely wide tables; fall back to key-value blocks.
        if len(keys) <= 6:
            table = doc.add_table(rows=1, cols=len(keys))
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_repeat_table_header(table.rows[0])
            for i, k in enumerate(keys):
                set_cell_text(table.rows[0].cells[i], humanize_key(k))
                set_cell_shading(table.rows[0].cells[i], "D9EAF7")
            for item in value:
                cells = table.add_row().cells
                for i, k in enumerate(keys):
                    set_cell_text(cells[i], stringify(item.get(k, "")))
            doc.add_paragraph("")
        else:
            for idx, item in enumerate(value, 1):
                doc.add_paragraph(f"[{idx}]", style=None).runs[0].bold = True
                add_kv_table(doc, [(humanize_key(k), v) for k, v in item.items()])
    else:
        add_list(doc, value)


def humanize_key(key: str) -> str:
    mapping = {
        "schema_version": "スキーマバージョン",
        "document": "文書管理",
        "source_documents": "参照元ドキュメント",
        "target_system": "対象システム",
        "design_policy": "設計方針",
        "traceability": "トレーサビリティ",
        "impacted_assets": "影響資産",
        "module_design": "モジュール設計",
        "data_design": "データ設計",
        "thread_design": "スレッド設計",
        "state_machine": "状態遷移設計",
        "processing_details": "処理詳細",
        "api_design": "API設計",
        "io_mapping": "I/Oマッピング",
        "consistency_and_concurrency": "整合性・排他設計",
        "error_safety_design": "エラー・安全設計",
        "performance_design": "性能設計",
        "build_and_test": "ビルド・テスト",
        "pre_review_self_check": "プレレビューセルフチェック",
        "test_policy": "テスト方針",
        "test_environment": "テスト環境",
        "test_data": "テストデータ",
        "test_cases": "テストケース",
        "coverage_matrix": "カバレッジマトリクス",
        "execution_procedure": "実行手順",
        "evidence_design": "証跡設計",
        "entry_exit_criteria": "開始・終了条件",
        "review_checklist": "レビューチェックリスト",
        "open_issues": "未確定事項",
    }
    return mapping.get(key, key.replace("_", " "))


def infer_doc_title(data: dict[str, Any]) -> str:
    candidates = [
        "document.document_name",
        "document_control.title",
        "feature.name",
        "title",
        "name",
    ]
    for path in candidates:
        v = get_path(data, path, default="")
        if isinstance(v, str) and v.strip():
            return v.strip()
    return "IBM-Bob 設計書"


def add_doc_control(doc: DocxDocument, data: dict[str, Any]) -> None:
    meta = get_path(data, "document", default={})
    if not isinstance(meta, dict):
        meta = get_path(data, "document_control", default={})
    if not isinstance(meta, dict):
        meta = {}
    rows = []
    for k in ["document_id", "document_name", "document_type", "version", "created_date", "author", "status"]:
        if k in meta:
            rows.append((humanize_key(k), meta[k]))
    if not rows:
        rows = [("生成日", str(date.today())), ("生成元", "IBM-Bob JSON")]
    add_kv_table(doc, rows, title="文書情報")


def build_fallback_document(data: dict[str, Any], out_path: str | Path, doc_type: str = "auto") -> None:
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    title = infer_doc_title(data)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)
    doc.add_paragraph("")
    add_doc_control(doc, data)

    top_order = ordered_top_keys(data)
    for key in top_order:
        if key in ("document", "document_control"):
            continue
        add_object_section(doc, key, data[key], level=1)
    doc.save(str(out_path))


def ordered_top_keys(data: dict[str, Any]) -> list[str]:
    preferred = [
        "schema_version", "source_documents", "requirements", "feature", "purpose_background", "scope", "system_context",
        "target_system", "design_policy", "traceability", "impacted_assets", "module_design", "data_design", "thread_design",
        "state_machine", "processing_details", "api_design", "io_mapping", "consistency_and_concurrency", "error_safety_design",
        "performance_design", "build_and_test", "test_policy", "test_environment", "test_data", "test_cases", "coverage_matrix",
        "execution_procedure", "evidence_design", "entry_exit_criteria", "review_checklist", "pre_review_self_check", "open_issues",
    ]
    keys = [k for k in preferred if k in data]
    keys.extend(k for k in data.keys() if k not in keys and k not in ("document", "document_control"))
    return keys


def extract_docx_text(docx_path: str | Path) -> str:
    import xml.etree.ElementTree as ET
    with zipfile.ZipFile(docx_path) as z:
        xml = z.read("word/document.xml")
    root = ET.fromstring(xml)
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    texts = [node.text or "" for node in root.findall(".//w:t", ns)]
    return "".join(texts)


def validate_output_basic(docx_path: str | Path, *, fail_on_unresolved: bool = True) -> list[str]:
    errors: list[str] = []
    path = Path(docx_path)
    if not path.exists() or path.stat().st_size == 0:
        return ["DOCX file does not exist or is empty."]
    try:
        with zipfile.ZipFile(path) as z:
            names = set(z.namelist())
            if "word/document.xml" not in names:
                errors.append("word/document.xml is missing.")
    except zipfile.BadZipFile:
        return ["Output is not a valid DOCX zip file."]
    try:
        text = extract_docx_text(path)
        if len(text.strip()) < 10:
            errors.append("Document text is too short; output may be empty.")
        if fail_on_unresolved:
            unresolved = sorted(set(UNRESOLVED_RE.findall(text)))
            if unresolved:
                preview = ", ".join(unresolved[:20])
                errors.append(f"Unresolved placeholders remain: {preview}")
    except Exception as e:
        errors.append(f"Failed to parse document.xml: {e}")
    return errors


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description="Generate IBM-Bob design DOCX from JSON and optional template.")
    parser.add_argument("--json", required=True, help="Input JSON file path")
    parser.add_argument("--template", help="Optional DOCX template path")
    parser.add_argument("--out", required=True, help="Output DOCX path")
    parser.add_argument("--doc-type", default="auto", choices=["auto", "basic", "detail", "function-test", "review"], help="Document type hint")
    parser.add_argument("--strict", action="store_true", help="Fail if unresolved placeholders remain")
    parser.add_argument("--fallback", action="store_true", help="Ignore template and generate a structured document from JSON")
    args = parser.parse_args(argv)

    data = load_json(args.json)
    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)

    used_template = False
    if args.template and not args.fallback:
        try:
            apply_template(data, args.template, out)
            used_template = True
        except Exception as e:
            print(f"[WARN] Template rendering failed: {e}", file=sys.stderr)
            print("[WARN] Falling back to structured DOCX generation.", file=sys.stderr)
            build_fallback_document(data, out, doc_type=args.doc_type)
    else:
        build_fallback_document(data, out, doc_type=args.doc_type)

    errors = validate_output_basic(out, fail_on_unresolved=args.strict)
    if errors:
        if used_template:
            print("[WARN] Generated DOCX did not pass validation. Falling back to structured DOCX generation.", file=sys.stderr)
            for err in errors:
                print(f"[WARN] {err}", file=sys.stderr)
            build_fallback_document(data, out, doc_type=args.doc_type)
            used_template = False
            errors = validate_output_basic(out, fail_on_unresolved=True)
        if errors:
            for err in errors:
                print(f"[ERROR] {err}", file=sys.stderr)
            return 2
    print(f"[OK] DOCX generated: {out}")
    print(f"[OK] mode: {'template' if used_template else 'fallback-structured'}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
