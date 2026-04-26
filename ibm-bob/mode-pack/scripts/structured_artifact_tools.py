# ibm-bob/mode-pack/scripts/structured_artifact_tools.py
# Provides small utilities for IBM-Bob structured specification artifacts.
# This exists so Bob modes can validate intermediate JSON, fill DOCX templates, and extract section snapshots.
# RELEVANT FILES: ibm-bob/mode-pack/structured-templates/**, ibm-bob/mode-pack/structured-artifact-workflow.md, ibm-bob/mode-pack/rules/ibmbob-sdlc-spec-reviewer/01-core.md
from __future__ import annotations

import argparse
import json
import re
import shutil
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception:  # pragma: no cover - allows list/validate even without python-docx
    Document = None  # type: ignore

from json_schema_lite import LiteSchemaValidator
from runtime_common import PACK_ROOT, ensure_dir

TEMPLATE_ROOT = PACK_ROOT / "structured-templates"

REGISTRY: dict[str, dict[str, str]] = {
    "basic_design": {
        "label": "基本設計書",
        "root": "basic-design",
        "template": "IBM-Bob_基本設計書テンプレート_v2.docx",
        "schema": "ibm_bob_basic_design_schema.json",
        "skeleton": "ibm_bob_basic_design_skeleton.json",
        "checklist": "ibm_bob_pre_review_checklist.json",
    },
    "detail_design": {
        "label": "詳細設計書",
        "root": "detail-design",
        "template": "IBM-Bob_詳細設計書テンプレート_v2.docx",
        "schema": "ibm_bob_detail_design_schema.json",
        "skeleton": "ibm_bob_detail_design_skeleton.json",
        "checklist": "ibm_bob_detail_design_pre_review_checklist.json",
    },
    "function_test_design": {
        "label": "機能テスト設計書",
        "root": "function-test-design",
        "template": "IBM-Bob_機能テスト設計書テンプレート_v1.docx",
        "schema": "ibm_bob_function_test_design_schema.json",
        "skeleton": "ibm_bob_function_test_design_skeleton.json",
        "checklist": "ibm_bob_function_test_pre_review_checklist.json",
    },
}

PLACEHOLDER_RE = re.compile(r"{{\s*json\.([^}]+?)\s*}}")
SECTION_RE = re.compile(r"\[SECTION_ID:\s*([^\]]+)\]")
JSON_PATH_RE = re.compile(r"JSON\s*Path\s*[:：]\s*([^\s]+)")


def _artifact_paths(kind: str) -> dict[str, Path]:
    if kind not in REGISTRY:
        raise SystemExit(f"unknown kind: {kind}. Use one of: {', '.join(REGISTRY)}")
    entry = REGISTRY[kind]
    root = TEMPLATE_ROOT / entry["root"]
    return {name: root / filename for name, filename in entry.items() if name in {"template", "schema", "skeleton", "checklist"}}


def _load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _write_json(path: Path, data: Any) -> None:
    ensure_dir(path.parent)
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def _json_lookup(data: Any, dotted_path: str) -> Any:
    current = data
    for part in dotted_path.strip().split("."):
        part = part.strip()
        if not part:
            continue
        if isinstance(current, dict):
            current = current.get(part, "")
        elif isinstance(current, list):
            if part.isdigit() and int(part) < len(current):
                current = current[int(part)]
            else:
                return ""
        else:
            return ""
    return current


def _to_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (str, int, float, bool)):
        return str(value)
    return json.dumps(value, ensure_ascii=False, indent=2)


def _replace_in_paragraph(paragraph, data: dict) -> None:
    text = paragraph.text
    if "{{" not in text:
        return
    replaced = PLACEHOLDER_RE.sub(lambda m: _to_text(_json_lookup(data, m.group(1))), text)
    if replaced == text:
        return
    # Preserve paragraph style but simplify runs. This is enough for generated drafts.
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = replaced
    else:
        paragraph.add_run(replaced)


def _replace_placeholders(doc, data: dict) -> int:
    count = 0
    all_paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paragraphs.extend(cell.paragraphs)
    for paragraph in all_paragraphs:
        before = paragraph.text
        _replace_in_paragraph(paragraph, data)
        if paragraph.text != before:
            count += 1
    return count


def command_list(_: argparse.Namespace) -> int:
    for kind, entry in REGISTRY.items():
        paths = _artifact_paths(kind)
        print(f"{kind}: {entry['label']}")
        for key in ("template", "schema", "skeleton", "checklist"):
            print(f"  {key}: {paths[key]}")
    return 0


def command_validate(args: argparse.Namespace) -> int:
    paths = _artifact_paths(args.kind)
    data = _load_json(Path(args.json))
    errors = LiteSchemaValidator().validate_file(data, paths["schema"])
    if errors:
        print("validation: NG")
        for error in errors:
            print(f"- {error}")
        return 1
    print("validation: OK")
    return 0


def command_new(args: argparse.Namespace) -> int:
    paths = _artifact_paths(args.kind)
    skeleton = _load_json(paths["skeleton"])
    _write_json(Path(args.out), skeleton)
    print(f"created: {args.out}")
    return 0


def command_render_docx(args: argparse.Namespace) -> int:
    if Document is None:
        raise SystemExit("python-docx is required for render-docx")
    paths = _artifact_paths(args.kind)
    data = _load_json(Path(args.json))
    errors = LiteSchemaValidator().validate_file(data, paths["schema"])
    if errors and not args.allow_invalid:
        print("validation: NG")
        for error in errors:
            print(f"- {error}")
        return 1
    out_path = Path(args.out)
    ensure_dir(out_path.parent)
    shutil.copy2(paths["template"], out_path)
    doc = Document(str(out_path))
    replaced_count = _replace_placeholders(doc, data)
    doc.save(str(out_path))
    print(f"created: {out_path}")
    print(f"placeholder_replacements: {replaced_count}")
    if errors:
        print("warning: output created with invalid JSON because --allow-invalid was supplied")
    return 0


def command_extract_sections(args: argparse.Namespace) -> int:
    if Document is None:
        raise SystemExit("python-docx is required for extract-sections")
    doc = Document(str(Path(args.docx)))
    sections: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None

    def push_text(text: str) -> None:
        nonlocal current
        stripped = text.strip()
        if not stripped:
            return
        section_match = SECTION_RE.search(stripped)
        if section_match:
            current = {"section_id": section_match.group(1), "json_path": "", "texts": []}
            sections.append(current)
        if current is not None:
            path_match = JSON_PATH_RE.search(stripped)
            if path_match and not current.get("json_path"):
                current["json_path"] = path_match.group(1).strip()
            current["texts"].append(stripped)

    for paragraph in doc.paragraphs:
        push_text(paragraph.text)
    for table_index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            sections.append({"section_id": f"TABLE-{table_index:03d}", "json_path": "", "table": rows})
    snapshot = {
        "artifact_kind": args.kind,
        "source_docx": str(Path(args.docx)),
        "extraction_policy": "section_id/json_path best-effort snapshot for IBM-Bob review and migration",
        "sections": sections,
    }
    _write_json(Path(args.out), snapshot)
    print(f"created: {args.out}")
    return 0


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="IBM-Bob structured artifact helper")
    sub = parser.add_subparsers(dest="command", required=True)

    p = sub.add_parser("list")
    p.set_defaults(func=command_list)

    p = sub.add_parser("validate")
    p.add_argument("--kind", required=True, choices=sorted(REGISTRY))
    p.add_argument("--json", required=True)
    p.set_defaults(func=command_validate)

    p = sub.add_parser("new")
    p.add_argument("--kind", required=True, choices=sorted(REGISTRY))
    p.add_argument("--out", required=True)
    p.set_defaults(func=command_new)

    p = sub.add_parser("render-docx")
    p.add_argument("--kind", required=True, choices=sorted(REGISTRY))
    p.add_argument("--json", required=True)
    p.add_argument("--out", required=True)
    p.add_argument("--allow-invalid", action="store_true")
    p.set_defaults(func=command_render_docx)

    p = sub.add_parser("extract-sections")
    p.add_argument("--kind", required=True, choices=sorted(REGISTRY))
    p.add_argument("--docx", required=True)
    p.add_argument("--out", required=True)
    p.set_defaults(func=command_extract_sections)
    return parser


def main() -> int:
    parser = build_parser()
    args = parser.parse_args()
    return args.func(args)


if __name__ == "__main__":
    raise SystemExit(main())
